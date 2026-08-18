"""
Resume Parser Module
====================
Extracts structured information from PDF and DOCX resumes.
Handles text extraction, cleaning, and entity recognition with graceful fallbacks.
"""

import re
import json
import logging
from pathlib import Path
from collections import Counter

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ─── Text Extraction ──────────────────────────────────────────────────────────

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file using PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF
        text = ""
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text("text") + "\n"
        doc.close()
        return text
    except ImportError:
        logger.warning("PyMuPDF not installed. Trying pdfplumber fallback.")
        return _extract_pdf_fallback(file_path)
    except Exception as e:
        logger.error(f"Error extracting PDF text: {e}")
        return ""


def _extract_pdf_fallback(file_path: str) -> str:
    """Fallback PDF extraction using pdfplumber if PyMuPDF unavailable."""
    try:
        import pdfplumber
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += (page.extract_text() or "") + "\n"
        return text
    except Exception as e:
        logger.error(f"PDF fallback extraction failed: {e}")
        return ""


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        import docx
        doc = docx.Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract table content
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error(f"Error extracting DOCX text: {e}")
        return ""


def clean_text(text: str) -> str:
    """Clean and normalize extracted text."""
    # Normalize line endings
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    # Remove excessive blank lines (keep max 2 consecutive newlines)
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Remove excessive spaces within lines
    lines = [re.sub(r'[ \t]+', ' ', line).strip() for line in text.split('\n')]
    return '\n'.join(lines).strip()


# ─── Entity Extraction ────────────────────────────────────────────────────────

def extract_email(text: str) -> str:
    """Extract email address using regex."""
    pattern = r'\b[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+\b'
    match = re.search(pattern, text)
    return match.group(0) if match else ""


def extract_phone(text: str) -> str:
    """Extract phone number using strict regex — avoids matching year ranges like 2019-23."""
    # Strict patterns that require actual phone number format
    patterns = [
        r'\+?\d{1,3}[-\s]?\(?\d{3}\)?[-\s]?\d{3}[-\s]?\d{4}',  # +91 9876543210
        r'\b[6-9]\d{9}\b',                                          # Indian mobile (starts 6-9)
        r'\b\d{3}[-.]\d{3}[-.]\d{4}\b',                           # 123-456-7890
        r'\(\d{3}\)\s?\d{3}[-.]\d{4}',                            # (123) 456-7890
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            val = match.group(0).strip()
            # Reject if it looks like a year range (e.g. 2019-23, 2019 - 2023)
            if re.match(r'^20\d{2}', val) or re.match(r'^19\d{2}', val):
                continue
            return val
    return ""


def extract_name(text: str, filename: str = "") -> str:
    """
    Extract candidate name using text-only strategies:
    1. 'Name: XXX' label in resume (most reliable for Indian resumes)
    2. spaCy NER PERSON entity
    3. Smart line scan — skips institutions, ALL-CAPS, long lines
    """

    # ── Strategy 1: Look for "Name:" label ────────────────────────────────────
    name_label_patterns = [
        r'(?:^|\n)\s*Name\s*[:]\s*([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+){1,4})',
        r'(?:^|\n)\s*Full\s*Name\s*[:]\s*([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+){1,4})',
        r'(?:^|\n)\s*Candidate\s*Name\s*[:]\s*([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+){1,4})',
        r'(?:^|\n)\s*Applicant\s*Name\s*[:]\s*([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+){1,4})',
    ]
    for pattern in name_label_patterns:
        match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
        if match:
            name = match.group(1).strip()
            if 2 <= len(name.split()) <= 5 and len(name) >= 4:
                return name

    # Words that indicate it's NOT a person's name
    NOT_NAME_KEYWORDS = [
        'COLLEGE', 'UNIVERSITY', 'INSTITUTE', 'SCHOOL', 'ACADEMY', 'POLYTECHNIC',
        'TECHNOLOGIES', 'SOLUTIONS', 'SYSTEMS', 'SERVICES', 'COMPANY', 'CORP',
        'LIMITED', 'LTD', 'PVT', 'INC', 'FOUNDATION', 'TRUST', 'CENTRE', 'CENTER',
        'RESUME', 'CV', 'CURRICULUM', 'PROFILE', 'SUMMARY', 'OBJECTIVE',
        'DEPARTMENT', 'FACULTY', 'MANAGEMENT', 'ENGINEERING', 'SCIENCE',
        'EDUCATION', 'EXPERIENCE', 'SKILLS', 'CONTACT', 'ADDRESS',
        'PHARMACY', 'MEDICAL', 'HOSPITAL', 'CLINIC', 'BANK', 'FINANCE',
        'NARAYANA', 'CHAITANYA', 'BHARATI', 'VIDYALAYA', 'VIKAS', 'KENDRIYA',
        'INTERNATIONAL', 'NATIONAL', 'GLOBAL', 'PUBLIC', 'PRIVATE', 'GOVERNMENT',
        'BHAVAN', 'NIKETAN', 'MANDIR', 'PEETH', 'SADAN', 'PARISHAD',
        'BOARD', 'COUNCIL', 'SOCIETY', 'COMMITTEE', 'HIGH', 'HIGHER', 'SECONDARY',
    ]

    # ── Strategy 2: spaCy NER ─────────────────────────────────────────────────
    try:
        import spacy
        nlp = spacy.load("en_core_web_sm")
        doc = nlp(text[:800])
        for ent in doc.ents:
            if ent.label_ == "PERSON" and len(ent.text.split()) >= 2:
                name = ent.text.strip()
                if not any(kw in name.upper() for kw in NOT_NAME_KEYWORDS):
                    return name
    except Exception:
        pass

    # ── Strategy 3: Smart line scan ───────────────────────────────────────────
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    for line in lines[:15]:
        # Skip if contains digits (year, phone, roll number etc)
        if re.search(r'\d', line):
            continue
        # Skip email / URL lines
        if '@' in line or 'www.' in line.lower() or 'http' in line.lower():
            continue
        # Skip lines with institution/company keywords
        if any(kw in line.upper() for kw in NOT_NAME_KEYWORDS):
            continue
        # Skip very long lines (descriptions, addresses)
        if len(line) > 45:
            continue
        # Skip lines that are ALL CAPS and more than 2 words (institution names)
        if line.isupper() and len(line.split()) > 2:
            continue
        # Skip single-word lines
        words = line.split()
        if len(words) < 2 or len(words) > 5:
            continue
        # Skip lines with punctuation like commas, colons (addresses, titles)
        if any(c in line for c in [',', ':', ';', '|', '/']):
            continue
        # Each word should start with capital letter (proper noun pattern)
        if all(w[0].isupper() for w in words if w):
            return line

    return ""






def extract_education(text: str) -> list:
    """Extract ONLY actual degree/qualification entries — skips objective, summary, seeking lines."""
    degree_keywords = [
        'B.Tech', 'B.E', 'B.Sc', 'B.Com', 'B.A', 'BCA', 'BBA',
        'M.Tech', 'M.E', 'M.Sc', 'MBA', 'MCA', 'M.A', 'M.Com',
        'Ph.D', 'PhD', 'Bachelor', 'Master', 'Doctorate',
        'B.S.', 'M.S.', 'B.Eng', 'M.Eng',
        '10th', '12th', 'HSC', 'SSC', 'Diploma', 'Associate',
        'Bachelor of Technology', 'Bachelor of Science', 'Bachelor of Commerce',
        'Master of Technology', 'Master of Science', 'Master of Business',
        'Secondary School', 'Higher Secondary',
    ]

    # Words that indicate it's NOT an education entry (objective/summary/description)
    SKIP_WORDS = [
        'seeking', 'looking for', 'motivated', 'passionate', 'dedicated',
        'objective', 'summary', 'career', 'opportunity', 'position', 'role',
        'graduate in', 'student of', 'pursuing',  # only skip if these appear alone
        'entry-level', 'fresher', 'aspiring',
    ]

    found = []
    lines = text.split('\n')
    for line in lines:
        line_stripped = line.strip()
        if not line_stripped or len(line_stripped) < 4:
            continue
        line_lower = line_stripped.lower()

        # Must contain a degree keyword
        if not any(kw.lower() in line_lower for kw in degree_keywords):
            continue

        # Skip if it's clearly an objective/summary sentence
        if any(skip in line_lower for skip in SKIP_WORDS):
            continue

        # Skip very long lines (likely descriptions, not degree titles)
        if len(line_stripped) > 120:
            continue

        found.append(line_stripped)

    # Deduplicate
    seen = set()
    unique = []
    for item in found:
        key = item.lower()[:50]
        if key not in seen:
            seen.add(key)
            unique.append(item)
    return unique[:6]


def extract_experience_years(text: str) -> float:
    """
    Extract ACTUAL work experience years from resume.
    - Uses explicit "X years of experience" mentions
    - Looks ONLY inside Work Experience section for date ranges
    - Returns 0 for freshers / students with no work history
    - Never calculates from education or other section years
    """
    # 1. Explicit mention: "2 years of experience", "3+ years exp"
    explicit_patterns = [
        r'(\d+(?:\.\d+)?)\+?\s*years?\s*(?:of\s*)?(?:work\s*)?(?:experience|exp)',
        r'(?:experience|exp)\s*(?:of\s*)?(\d+(?:\.\d+)?)\+?\s*years?',
        r'(\d+(?:\.\d+)?)\+?\s*years?\s*in\s+(?:the\s+)?(?:industry|field|domain)',
    ]
    for pattern in explicit_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            val = float(match.group(1))
            if 0 < val < 40:  # sanity check
                return val

    # 2. Fresher / no experience keywords → return 0
    fresher_keywords = [
        'fresher', 'no experience', '0 years', 'zero experience',
        'no work experience', 'seeking entry', 'entry-level',
        'recent graduate', 'fresh graduate',
    ]
    text_lower = text.lower()
    if any(kw in text_lower for kw in fresher_keywords):
        return 0.0

    # 3. Look ONLY in Work Experience section for date ranges
    EXPERIENCE_SECTION_HEADERS = [
        'work experience', 'experience', 'employment history',
        'professional experience', 'internship', 'internships',
        'work history', 'career history', 'job experience',
    ]
    EDUCATION_SECTION_HEADERS = [
        'education', 'academic', 'qualification', 'degree',
        'certifications', 'projects', 'skills', 'declaration',
    ]

    # Extract just the experience section
    exp_section = ""
    text_lines = text.split('\n')
    in_exp_section = False
    for line in text_lines:
        line_lower = line.strip().lower()
        if any(h in line_lower for h in EXPERIENCE_SECTION_HEADERS) and len(line.strip()) < 40:
            in_exp_section = True
            continue
        if in_exp_section:
            if any(h in line_lower for h in EDUCATION_SECTION_HEADERS) and len(line.strip()) < 40:
                break
            exp_section += line + '\n'

    # Only compute from years found in the experience section
    if exp_section.strip():
        year_pattern = r'\b(20\d{2}|19\d{2})\b'
        years_found = [int(y) for y in re.findall(year_pattern, exp_section)]
        if len(years_found) >= 2:
            min_year = min(years_found)
            max_year = max(years_found)
            computed = max_year - min_year
            if 0 < computed < 30:
                return float(round(computed, 1))

    # 4. No experience found → 0 (student/fresher)
    return 0.0




def extract_section(text: str, section_names: list, next_sections: list = None) -> str:
    """
    Extract a named section from resume text.
    Returns text between the section header and the next section header.
    """
    if next_sections is None:
        next_sections = [
            'experience', 'education', 'skills', 'projects', 'certifications',
            'awards', 'publications', 'languages', 'interests', 'references',
            'work', 'employment', 'summary', 'objective', 'contact'
        ]

    # Build pattern for section headers
    section_pattern = '|'.join(re.escape(s) for s in section_names)
    next_section_pattern = '|'.join(re.escape(s) for s in next_sections)

    # Match section header (case-insensitive, at start of line or standalone)
    match = re.search(
        rf'(?i)^[\s\-_•]*({section_pattern})[\s\-_:•]*$',
        text, re.MULTILINE
    )
    if not match:
        return ""

    start = match.end()
    # Find next section
    next_match = re.search(
        rf'(?i)^[\s\-_•]*({next_section_pattern})[\s\-_:•]*$',
        text[start:], re.MULTILINE
    )
    end = start + next_match.start() if next_match else len(text)
    return text[start:end].strip()


def extract_projects(text: str) -> list:
    """Extract project titles/descriptions from the Projects section."""
    section_text = extract_section(text, ['projects', 'academic projects', 'personal projects', 'key projects'])
    if not section_text:
        return []

    projects = []
    lines = [line.strip() for line in section_text.split('\n') if line.strip()]
    for line in lines:
        # Skip bullet points and short lines that are likely descriptions
        clean_line = re.sub(r'^[\•\-\*\u2022\u25cf]\s*', '', line).strip()
        if len(clean_line) > 10:
            projects.append(clean_line)

    return projects[:10]  # Limit to 10 projects


def extract_certifications(text: str) -> list:
    """Extract certifications — stops at declaration/signature, skips non-cert content."""

    # Words that mark end of certifications (declaration section)
    STOP_WORDS = [
        'declaration', 'i hereby declare', 'i declare', 'signature',
        'place:', 'date:', 'references', 'referee',
        'hereby declare that', 'information provided above',
        'true and correct', 'knowledge and belief',
    ]

    # Words that indicate a line is NOT a certification
    INVALID_CERT_WORDS = [
        'declaration', 'signature', 'hereby', 'declare', 'knowledge',
        'belief', 'true and correct', 'information provided',
        'place', 'date', 'sincerely', 'regards', 'thank you',
        'linkedin', 'github', 'http', 'www.', '.com', '.in', '@',
    ]

    # Try section-based extraction first
    section_text = extract_section(
        text,
        ['certifications', 'certificates', 'certification', 'credentials', 'courses', 'achievements']
    )

    certifications = []
    if section_text:
        lines = [line.strip() for line in section_text.split('\n') if line.strip()]
        for line in lines:
            line_lower = line.lower()

            # Stop if we hit a declaration/signature block
            if any(stop in line_lower for stop in STOP_WORDS):
                break

            # Skip lines with invalid content
            if any(inv in line_lower for inv in INVALID_CERT_WORDS):
                continue

            clean_line = re.sub(r'^[\•\-\*\u2022\u25cf\d\.]+\s*', '', line).strip()

            # Skip very short or very long lines
            if len(clean_line) < 6 or len(clean_line) > 200:
                continue

            certifications.append(clean_line)

    # Also look for well-known certification patterns inline (only if section not found)
    if not certifications:
        cert_patterns = [
            r'(?:NPTEL|Coursera|edX|Udemy|LinkedIn Learning)[\s:]+[\w\s\-&()]{5,80}',
            r'AWS\s+Certified\s+[\w\s]{3,50}',
            r'Google\s+(?:Cloud\s+)?Certified\s+[\w\s]{3,50}',
            r'Microsoft\s+Certified[\s:]+[\w\s]{3,50}',
            r'Cisco\s+Certified\s+[\w\s]{3,50}',
        ]
        for pattern in cert_patterns:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                cert = match.group(0).strip()
                if cert not in certifications and len(cert) > 8:
                    certifications.append(cert)

    return certifications[:8]


def extract_keywords(text: str, top_n: int = 20) -> list:
    """
    Extract top keywords using word frequency analysis with stopword removal.
    Returns most meaningful/frequent technical terms from the resume.
    """
    STOPWORDS = {
        'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
        'of', 'with', 'by', 'from', 'up', 'about', 'into', 'through', 'during',
        'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had',
        'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may', 'might',
        'shall', 'can', 'not', 'no', 'nor', 'so', 'yet', 'both', 'either',
        'i', 'me', 'my', 'myself', 'we', 'our', 'you', 'your', 'he', 'she',
        'they', 'them', 'this', 'that', 'these', 'those', 'it', 'its',
        'as', 'if', 'then', 'than', 'while', 'when', 'where', 'how', 'who',
        'which', 'what', 'also', 'such', 'more', 'most', 'other', 'some', 'any',
        'all', 'each', 'every', 'both', 'few', 'more', 'most',
        'work', 'worked', 'working', 'developed', 'development', 'using', 'used',
        'experience', 'experienced', 'knowledge', 'well', 'good', 'strong',
        'ability', 'skills', 'skill', 'team', 'responsible', 'responsibilities',
        'year', 'years', 'month', 'months', 'new', 'current', 'previous',
    }

    # Extract words (2+ chars, alphabetic)
    words = re.findall(r'\b[a-zA-Z][a-zA-Z+#.]{1,}\b', text)
    # Filter stopwords and short words
    meaningful = [w.lower() for w in words if w.lower() not in STOPWORDS and len(w) >= 3]
    # Count frequencies
    counter = Counter(meaningful)
    # Return top N by frequency
    top_keywords = [word for word, count in counter.most_common(top_n) if count >= 2]
    return top_keywords


def compute_resume_score(parsed_data: dict) -> float:
    """
    Compute a resume completeness/quality score (0-100) based on:
    - Has name, email, phone
    - Has education info
    - Has extracted skills
    - Has experience years
    - Has projects
    - Has certifications
    - Text length (content density)
    """
    score = 0.0

    # Contact info completeness (20 points)
    if parsed_data.get("name") and parsed_data["name"] != "Unknown":
        score += 8
    if parsed_data.get("email"):
        score += 7
    if parsed_data.get("phone"):
        score += 5

    # Education (15 points)
    education = parsed_data.get("education", [])
    if isinstance(education, str):
        education = json.loads(education) if education else []
    if education:
        score += 15

    # Skills (25 points)
    skills = parsed_data.get("skills", [])
    skill_count = len(skills)
    if skill_count >= 10:
        score += 25
    elif skill_count >= 5:
        score += 15
    elif skill_count >= 2:
        score += 8

    # Experience (20 points)
    exp_years = parsed_data.get("experience_years", 0)
    if exp_years >= 3:
        score += 20
    elif exp_years >= 1:
        score += 12
    elif exp_years > 0:
        score += 5

    # Projects (10 points)
    projects = parsed_data.get("projects", [])
    if isinstance(projects, str):
        projects = json.loads(projects) if projects else []
    if len(projects) >= 3:
        score += 10
    elif len(projects) >= 1:
        score += 5

    # Certifications (5 points)
    certs = parsed_data.get("certifications", [])
    if isinstance(certs, str):
        certs = json.loads(certs) if certs else []
    if certs:
        score += 5

    # Text richness (5 points)
    text = parsed_data.get("text", "")
    if len(text) > 2000:
        score += 5
    elif len(text) > 500:
        score += 2

    return min(100.0, round(score, 1))


# ─── Main Parser ──────────────────────────────────────────────────────────────

def parse_resume(file_path: str) -> dict:
    """
    Main resume parser function.
    Uses Groq AI for intelligent extraction when available,
    falls back to regex-based extraction if Groq is not configured.

    Returns:
        dict with keys: text, name, email, phone, education, experience_years,
                         projects, certifications, keywords, summary
    """
    file_path_obj = Path(file_path)
    ext = file_path_obj.suffix.lower()

    # Extract raw text based on file type
    raw_text = ""
    if ext == '.pdf':
        raw_text = extract_text_from_pdf(file_path)
    elif ext in ('.docx', '.doc'):
        raw_text = extract_text_from_docx(file_path)
    else:
        logger.error(f"Unsupported file type: {ext}")

    if not raw_text:
        logger.warning(f"No text extracted from {file_path}")

    cleaned_text = clean_text(raw_text)

    # ── Try Groq AI parsing first ──────────────────────────────────────────
    try:
        from app.ai.groq_engine import parse_resume_with_groq, is_groq_available
        if is_groq_available():
            logger.info("Using Groq AI for resume parsing...")
            groq_result = parse_resume_with_groq(cleaned_text)
            if groq_result:
                logger.info(f"Groq parsed: {groq_result.get('name')} | {len(groq_result.get('skills', []))} skills")
                return {
                    "text": cleaned_text,
                    "name": groq_result.get("name") or extract_name(raw_text, file_path_obj.name),
                    "email": groq_result.get("email") or extract_email(cleaned_text),
                    "phone": groq_result.get("phone") or extract_phone(cleaned_text),
                    "education": json.dumps(groq_result.get("education") or extract_education(raw_text)),
                    "experience_years": float(groq_result.get("experience_years") or 0),
                    "projects": json.dumps(groq_result.get("projects") or extract_projects(raw_text)),
                    "certifications": json.dumps(groq_result.get("certifications") or extract_certifications(raw_text)),
                    "keywords": json.dumps(groq_result.get("keywords") or extract_keywords(cleaned_text)),
                    "summary": groq_result.get("summary", ""),
                    # Also return skills so router can use them
                    "_groq_skills": groq_result.get("skills", []),
                    "_parsed_by": "groq",
                }
    except Exception as e:
        logger.warning(f"Groq parsing failed, falling back to regex: {e}")

    # ── Regex fallback ─────────────────────────────────────────────────────
    logger.info("Using regex-based resume parsing...")
    education_list  = extract_education(raw_text)
    projects_list   = extract_projects(raw_text)
    certs_list      = extract_certifications(raw_text)
    keywords_list   = extract_keywords(cleaned_text)

    return {
        "text": cleaned_text,
        "name": extract_name(raw_text, file_path_obj.name),
        "email": extract_email(cleaned_text),
        "phone": extract_phone(cleaned_text),
        "education": json.dumps(education_list),
        "experience_years": extract_experience_years(cleaned_text),
        "projects": json.dumps(projects_list),
        "certifications": json.dumps(certs_list),
        "keywords": json.dumps(keywords_list),
        "summary": "",
        "_parsed_by": "regex",
    }
